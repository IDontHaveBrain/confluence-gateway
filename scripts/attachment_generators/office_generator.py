"""Office document (DOCX, XLSX) attachment generator"""

import io
import random
from typing import Dict, Any
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.chart import BarChart, Reference, LineChart
from openpyxl.utils import get_column_letter

from .metadata import ATTACHMENT_METADATA

class OfficeGenerator:
    """Generates Office document attachments (DOCX, XLSX)"""
    
    def __init__(self, config, file_type: str = "docx"):
        self.config = config
        self.file_type = file_type
    
    def generate(self, category: str) -> Dict[str, Any]:
        """Generate Office document based on file type and category"""
        if self.file_type == "docx":
            return self._generate_docx(category)
        elif self.file_type == "xlsx":
            return self._generate_xlsx(category)
        else:
            raise ValueError(f"Unsupported file type: {self.file_type}")
    
    def _generate_docx(self, category: str) -> Dict[str, Any]:
        """Generate DOCX document"""
        doc = Document()
        
        generators = {
            "api": self._create_api_tutorial,
            "technical": self._create_installation_guide,
            "project": self._create_meeting_minutes,
            "planning": self._create_meeting_minutes,
            "releases": self._create_status_report
        }
        
        generator_func = generators.get(category, self._create_installation_guide)
        generator_func(doc)

        category_key = category if category in ["api", "technical", "project"] else "technical"
        filenames = ATTACHMENT_METADATA["docx"].get(category_key, ["Document.docx"])
        filename = f"{self.config.prefix}_{random.choice(filenames)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return {
            "filename": filename,
            "content": buffer.getvalue(),
            "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "size": buffer.tell()
        }
    
    def _generate_xlsx(self, category: str) -> Dict[str, Any]:
        """Generate XLSX spreadsheet"""
        # Create workbook
        wb = Workbook()
        
        generators = {
            "api": self._create_api_matrix,
            "technical": self._create_performance_metrics,
            "project": self._create_project_timeline,
            "planning": self._create_project_timeline,
            "releases": self._create_resource_allocation
        }
        
        generator_func = generators.get(category, self._create_performance_metrics)
        generator_func(wb)

        category_key = category if category in ["api", "technical", "project"] else "technical"
        filenames = ATTACHMENT_METADATA["xlsx"].get(category_key, ["Spreadsheet.xlsx"])
        filename = f"{self.config.prefix}_{random.choice(filenames)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        return {
            "filename": filename,
            "content": buffer.getvalue(),
            "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "size": buffer.tell()
        }
    
    # DOCX Content Generators
    
    def _create_api_tutorial(self, doc: Document) -> None:
        """Create API tutorial document"""

        title = doc.add_heading('API Integration Tutorial', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(
            'This tutorial will guide you through the process of integrating with our REST API. '
            'You will learn how to authenticate, make requests, and handle responses effectively.'
        )
        
        # Prerequisites
        doc.add_heading('Prerequisites', level=1)
        doc.add_paragraph('Before starting this tutorial, ensure you have:')
        for item in ['API credentials (client ID and secret)', 'A development environment', 'Basic knowledge of HTTP and REST']:
            doc.add_paragraph(item, style='List Bullet')
        
        # Authentication section
        doc.add_heading('Authentication', level=1)
        doc.add_paragraph('All API requests require authentication using OAuth 2.0. Follow these steps:')
        
        # Code example
        doc.add_heading('Step 1: Obtain Access Token', level=2)
        code_para = doc.add_paragraph()
        code_para.add_run('POST /oauth/token\n').font.name = 'Courier New'
        code_para.add_run('Content-Type: application/json\n\n').font.name = 'Courier New'
        code_para.add_run('{\n  "client_id": "your_client_id",\n  "client_secret": "your_secret",\n  "grant_type": "client_credentials"\n}').font.name = 'Courier New'
        
        # Making requests
        doc.add_heading('Making API Requests', level=1)
        doc.add_paragraph(
            'Once authenticated, include the access token in the Authorization header of all requests:'
        )
        
        example = doc.add_paragraph()
        example.add_run('Authorization: Bearer YOUR_ACCESS_TOKEN').font.name = 'Courier New'
        
        # Example requests
        doc.add_heading('Example: List Resources', level=2)
        doc.add_paragraph('To retrieve a list of resources:')
        
        request = doc.add_paragraph()
        request.add_run('GET /api/v1/resources?limit=10&page=1').font.name = 'Courier New'
        
        # Error handling
        doc.add_heading('Error Handling', level=1)
        doc.add_paragraph('The API uses standard HTTP status codes. Common errors include:')
        
        # Error table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Status Code'
        hdr_cells[1].text = 'Error Type'
        hdr_cells[2].text = 'Description'
        
        errors = [
            ('400', 'Bad Request', 'Invalid request parameters'),
            ('401', 'Unauthorized', 'Missing or invalid authentication'),
            ('403', 'Forbidden', 'Insufficient permissions'),
            ('404', 'Not Found', 'Resource does not exist'),
            ('429', 'Too Many Requests', 'Rate limit exceeded')
        ]
        
        for code, error, desc in errors:
            row_cells = table.add_row().cells
            row_cells[0].text = code
            row_cells[1].text = error
            row_cells[2].text = desc
    
    def _create_installation_guide(self, doc: Document) -> None:
        """Create installation guide document"""

        doc.add_heading('Installation Guide', 0)
        
        # System requirements
        doc.add_heading('System Requirements', level=1)
        doc.add_paragraph('Ensure your system meets the following requirements:')
        
        req_table = doc.add_table(rows=1, cols=2)
        req_table.style = 'Light List Accent 1'
        hdr_cells = req_table.rows[0].cells
        hdr_cells[0].text = 'Component'
        hdr_cells[1].text = 'Requirement'
        
        requirements = [
            ('Operating System', 'Ubuntu 20.04 LTS or later'),
            ('CPU', 'Minimum 4 cores, 8 cores recommended'),
            ('Memory', 'Minimum 8GB RAM, 16GB recommended'),
            ('Storage', '50GB available disk space'),
            ('Network', 'Stable internet connection')
        ]
        
        for component, req in requirements:
            row_cells = req_table.add_row().cells
            row_cells[0].text = component
            row_cells[1].text = req
        
        doc.add_page_break()
        
        # Installation steps
        doc.add_heading('Installation Steps', level=1)
        
        steps = [
            ('Update System Packages', 'sudo apt-get update && sudo apt-get upgrade -y'),
            ('Install Dependencies', 'sudo apt-get install -y docker docker-compose git'),
            ('Clone Repository', 'git clone https://github.com/example/application.git'),
            ('Configure Environment', 'cp .env.example .env && nano .env'),
            ('Start Services', 'docker-compose up -d'),
            ('Verify Installation', 'docker-compose ps')
        ]
        
        for i, (step_title, command) in enumerate(steps, 1):
            doc.add_heading(f'Step {i}: {step_title}', level=2)
            doc.add_paragraph('Run the following command:')
            cmd_para = doc.add_paragraph()
            cmd_para.add_run(command).font.name = 'Courier New'
        
        # Post-installation
        doc.add_heading('Post-Installation Configuration', level=1)
        doc.add_paragraph(
            'After successful installation, configure the application by editing the configuration '
            'file located at /etc/application/config.yaml. Key settings include:'
        )
        
        for setting in ['Database connection strings', 'API endpoints', 'Security certificates', 'Logging levels']:
            doc.add_paragraph(f'• {setting}')
    
    def _create_meeting_minutes(self, doc: Document) -> None:
        """Create meeting minutes document"""
        # Header
        doc.add_heading('Meeting Minutes', 0)
        
        # Meeting info
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Light Shading Accent 1'
        
        meeting_date = datetime.now().strftime('%B %d, %Y')
        info_data = [
            ('Date:', meeting_date),
            ('Time:', '2:00 PM - 3:30 PM'),
            ('Location:', 'Conference Room A / Virtual'),
            ('Attendees:', 'John Doe, Jane Smith, Bob Johnson, Alice Brown')
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
        
        # Agenda
        doc.add_heading('Agenda', level=1)
        agenda_items = [
            'Sprint Review and Retrospective',
            'Technical Architecture Discussion',
            'Resource Allocation for Q2',
            'Risk Assessment Update',
            'Action Items Review'
        ]
        
        for item in agenda_items:
            doc.add_paragraph(item, style='List Number')
        
        # Discussion points
        doc.add_heading('Discussion Points', level=1)
        
        discussions = [
            ('Sprint Performance', 'The team successfully completed 85% of planned story points. '
             'Velocity has improved by 15% compared to the previous sprint.'),
            ('Technical Debt', 'Identified areas requiring refactoring in the authentication module. '
             'Proposed dedicating 20% of next sprint to address technical debt.'),
            ('New Requirements', 'Product owner presented new requirements for the reporting feature. '
             'Team estimated 3 sprints for full implementation.')
        ]
        
        for topic, content in discussions:
            doc.add_heading(topic, level=2)
            doc.add_paragraph(content)
        
        # Action items
        doc.add_heading('Action Items', level=1)
        
        action_table = doc.add_table(rows=1, cols=4)
        action_table.style = 'Light Grid Accent 1'
        hdr_cells = action_table.rows[0].cells
        headers = ['Action Item', 'Owner', 'Due Date', 'Priority']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
        
        actions = [
            ('Update project roadmap', 'John Doe', 'March 31', 'High'),
            ('Schedule architecture review', 'Jane Smith', 'April 5', 'Medium'),
            ('Prepare budget proposal', 'Bob Johnson', 'April 10', 'High'),
            ('Document API changes', 'Alice Brown', 'March 28', 'Medium')
        ]
        
        for action, owner, due, priority in actions:
            row_cells = action_table.add_row().cells
            row_cells[0].text = action
            row_cells[1].text = owner
            row_cells[2].text = due
            row_cells[3].text = priority
    
    def _create_status_report(self, doc: Document) -> None:
        """Create project status report"""

        doc.add_heading('Project Status Report', 0)
        doc.add_paragraph(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
        
        # Executive summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(
            'The project is progressing according to schedule with 75% of Phase 2 milestones completed. '
            'Current sprint velocity is meeting targets, and the team has successfully resolved '
            'all critical blockers from the previous reporting period.'
        )
        
        # Status overview
        doc.add_heading('Overall Status', level=1)
        
        status_table = doc.add_table(rows=1, cols=3)
        status_table.style = 'Light List Accent 1'
        hdr_cells = status_table.rows[0].cells
        hdr_cells[0].text = 'Area'
        hdr_cells[1].text = 'Status'
        hdr_cells[2].text = 'Notes'
        
        status_data = [
            ('Schedule', 'On Track', 'Currently in Sprint 12 of 16'),
            ('Budget', 'Under Budget', '85% of budget consumed, 90% work complete'),
            ('Scope', 'On Track', 'All planned features in development'),
            ('Quality', 'Exceeding', 'Defect rate below target threshold'),
            ('Resources', 'Adequate', 'Full team allocation through Q2')
        ]
        
        for area, status, notes in status_data:
            row_cells = status_table.add_row().cells
            row_cells[0].text = area
            row_cells[1].text = status
            row_cells[2].text = notes
        
        # Key accomplishments
        doc.add_heading('Key Accomplishments', level=1)
        accomplishments = [
            'Completed user authentication module with OAuth2 integration',
            'Successfully migrated database to new infrastructure',
            'Achieved 95% code coverage for core modules',
            'Deployed beta version to staging environment',
            'Completed security audit with no critical findings'
        ]
        
        for item in accomplishments:
            doc.add_paragraph(f'✓ {item}')
    
    # XLSX Content Generators
    
    def _create_api_matrix(self, wb: Workbook) -> None:
        """Create API endpoints matrix spreadsheet"""
        ws = wb.active
        ws.title = "API Endpoints"
        
        # Headers
        headers = ['Endpoint', 'Method', 'Description', 'Auth Required', 'Rate Limit', 'Response Time (ms)']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")
        
        # Data
        endpoints = [
            ('/api/v1/users', 'GET', 'List all users', 'Yes', '100/hour', 150),
            ('/api/v1/users', 'POST', 'Create new user', 'Yes', '50/hour', 200),
            ('/api/v1/users/{id}', 'GET', 'Get user details', 'Yes', '200/hour', 100),
            ('/api/v1/users/{id}', 'PUT', 'Update user', 'Yes', '100/hour', 180),
            ('/api/v1/users/{id}', 'DELETE', 'Delete user', 'Yes', '50/hour', 150),
            ('/api/v1/auth/login', 'POST', 'User login', 'No', '20/hour', 300),
            ('/api/v1/auth/refresh', 'POST', 'Refresh token', 'Yes', '100/hour', 120),
            ('/api/v1/products', 'GET', 'List products', 'No', '500/hour', 180),
            ('/api/v1/orders', 'GET', 'List orders', 'Yes', '200/hour', 250),
            ('/api/v1/orders', 'POST', 'Create order', 'Yes', '100/hour', 350),
        ]
        
        for row, data in enumerate(endpoints, 2):
            for col, value in enumerate(data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                if col == 4:  # Auth required column
                    cell.alignment = Alignment(horizontal="center")
                if col == 6:  # Response time column
                    if value > 200:
                        cell.font = Font(color="FF0000")  # Red for slow responses
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Add summary sheet
        summary_ws = wb.create_sheet("Summary")
        summary_ws['A1'] = "API Statistics"
        summary_ws['A1'].font = Font(bold=True, size=14)
        
        stats = [
            ['Total Endpoints', len(endpoints)],
            ['Authenticated Endpoints', sum(1 for e in endpoints if e[3] == 'Yes')],
            ['Public Endpoints', sum(1 for e in endpoints if e[3] == 'No')],
            ['Average Response Time', f"{sum(e[5] for e in endpoints) / len(endpoints):.0f} ms"]
        ]
        
        for row, (label, value) in enumerate(stats, 3):
            summary_ws.cell(row=row, column=1, value=label).font = Font(bold=True)
            summary_ws.cell(row=row, column=2, value=value)
    
    def _create_performance_metrics(self, wb: Workbook) -> None:
        """Create performance metrics spreadsheet"""
        ws = wb.active
        ws.title = "Performance Metrics"
        
        # Headers
        headers = ['Date', 'Response Time (ms)', 'Throughput (req/s)', 'Error Rate (%)', 
                  'CPU Usage (%)', 'Memory Usage (GB)', 'Active Connections']
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        
        # Generate sample data for last 30 days
        base_date = datetime.now() - timedelta(days=30)
        for day in range(30):
            date = base_date + timedelta(days=day)
            row = day + 2
            
            ws.cell(row=row, column=1, value=date.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=2, value=random.randint(80, 150))  # Response time
            ws.cell(row=row, column=3, value=random.randint(800, 1200))  # Throughput
            ws.cell(row=row, column=4, value=round(random.uniform(0.1, 2.0), 2))  # Error rate
            ws.cell(row=row, column=5, value=random.randint(45, 85))  # CPU usage
            ws.cell(row=row, column=6, value=round(random.uniform(2.5, 4.5), 1))  # Memory
            ws.cell(row=row, column=7, value=random.randint(200, 500))  # Connections
        
        # Add chart for response time
        chart = LineChart()
        chart.title = "Response Time Trend"
        chart.style = 13
        chart.y_axis.title = 'Response Time (ms)'
        chart.x_axis.title = 'Date'
        
        data = Reference(ws, min_col=2, min_row=1, max_row=31)
        dates = Reference(ws, min_col=1, min_row=2, max_row=31)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(dates)
        
        ws.add_chart(chart, "I2")
        
        # Add conditional formatting for error rate
        from openpyxl.formatting.rule import CellIsRule
        
        red_fill = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
        ws.conditional_formatting.add('D2:D31',
            CellIsRule(operator='greaterThan', formula=['1.5'], fill=red_fill))
    
    def _create_project_timeline(self, wb: Workbook) -> None:
        """Create project timeline spreadsheet"""
        ws = wb.active
        ws.title = "Project Timeline"
        
        # Headers
        headers = ['Task', 'Start Date', 'End Date', 'Duration (days)', 'Status', 
                  'Assigned To', 'Progress (%)', 'Dependencies']
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="70AD47", end_color="70AD47", fill_type="solid")
        
        # Project tasks
        tasks = [
            ('Project Kickoff', 0, 1, 'Complete', 'PM Team', 100, ''),
            ('Requirements Gathering', 1, 10, 'Complete', 'BA Team', 100, 'Task 1'),
            ('System Design', 8, 15, 'Complete', 'Architects', 100, 'Task 2'),
            ('Database Design', 12, 10, 'In Progress', 'DB Team', 80, 'Task 3'),
            ('API Development', 15, 20, 'In Progress', 'Backend Team', 60, 'Task 3'),
            ('Frontend Development', 18, 25, 'In Progress', 'Frontend Team', 40, 'Task 3'),
            ('Integration Testing', 35, 10, 'Not Started', 'QA Team', 0, 'Task 5,6'),
            ('User Acceptance Testing', 40, 5, 'Not Started', 'QA Team', 0, 'Task 7'),
            ('Deployment Preparation', 42, 3, 'Not Started', 'DevOps', 0, 'Task 8'),
            ('Production Release', 45, 1, 'Not Started', 'All Teams', 0, 'Task 9'),
        ]
        
        base_date = datetime.now() - timedelta(days=20)
        
        for row, (task, start_offset, duration, status, assigned, progress, deps) in enumerate(tasks, 2):
            start_date = base_date + timedelta(days=start_offset)
            end_date = start_date + timedelta(days=duration)
            
            ws.cell(row=row, column=1, value=task)
            ws.cell(row=row, column=2, value=start_date.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=3, value=end_date.strftime('%Y-%m-%d'))
            ws.cell(row=row, column=4, value=duration)
            ws.cell(row=row, column=5, value=status)
            ws.cell(row=row, column=6, value=assigned)
            ws.cell(row=row, column=7, value=progress)
            ws.cell(row=row, column=8, value=deps)
            
            # Color code status
            status_cell = ws.cell(row=row, column=5)
            if status == 'Complete':
                status_cell.fill = PatternFill(start_color='C5E0B4', end_color='C5E0B4', fill_type='solid')
            elif status == 'In Progress':
                status_cell.fill = PatternFill(start_color='FFE699', end_color='FFE699', fill_type='solid')
            else:
                status_cell.fill = PatternFill(start_color='F8CBAD', end_color='F8CBAD', fill_type='solid')
        
        # Auto-adjust columns
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
    
    def _create_resource_allocation(self, wb: Workbook) -> None:
        """Create resource allocation spreadsheet"""
        ws = wb.active
        ws.title = "Resource Allocation"
        
        # Headers
        headers = ['Resource', 'Role', 'Current Project', 'Allocation (%)', 
                  'Available From', 'Hourly Rate', 'Monthly Cost']
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="ED7D31", end_color="ED7D31", fill_type="solid")
        
        # Resources
        resources = [
            ('John Doe', 'Senior Developer', 'Project Alpha', 100, 'N/A', 150, 24000),
            ('Jane Smith', 'Tech Lead', 'Project Alpha', 80, 'April 1', 180, 23040),
            ('Bob Johnson', 'Frontend Dev', 'Project Beta', 60, 'March 25', 120, 11520),
            ('Alice Brown', 'QA Engineer', 'Project Alpha', 100, 'N/A', 100, 16000),
            ('Charlie Davis', 'DevOps Engineer', 'Infrastructure', 40, 'Available', 140, 8960),
            ('Eve Wilson', 'Business Analyst', 'Project Alpha', 100, 'N/A', 110, 17600),
            ('Frank Miller', 'UI/UX Designer', 'Project Beta', 50, 'April 15', 130, 10400),
            ('Grace Lee', 'Database Admin', 'Project Alpha', 70, 'Available', 140, 15680),
        ]
        
        for row, data in enumerate(resources, 2):
            for col, value in enumerate(data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                
                # Format currency columns
                if col in [6, 7]:
                    cell.number_format = '$#,##0'
                
                # Color code allocation
                if col == 4:
                    if value >= 100:
                        cell.fill = PatternFill(start_color='F8CBAD', end_color='F8CBAD', fill_type='solid')
                    elif value >= 80:
                        cell.fill = PatternFill(start_color='FFE699', end_color='FFE699', fill_type='solid')
        
        # Add summary row
        summary_row = len(resources) + 3
        ws.cell(row=summary_row, column=1, value='TOTAL').font = Font(bold=True)
        ws.cell(row=summary_row, column=7, value=f'=SUM(G2:G{len(resources)+1})').font = Font(bold=True)
        
        # Add allocation chart
        chart = BarChart()
        chart.title = "Resource Allocation"
        chart.style = 10
        chart.y_axis.title = 'Allocation %'
        
        data = Reference(ws, min_col=4, min_row=1, max_row=len(resources)+1)
        categories = Reference(ws, min_col=1, min_row=2, max_row=len(resources)+1)
        chart.add_data(data, titles_from_data=True)
        chart.set_categories(categories)
        
        ws.add_chart(chart, "I2")