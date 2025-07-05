#!/usr/bin/env python3
"""
Attachment Generators for Real Data

This module handles attachment generation using real or sample files
for testing purposes.
"""

import io
import json
import logging
import random
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any

from PIL import Image
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from openpyxl import Workbook

from confluence_gateway.adapters.confluence.client import ConfluenceClient

logger = logging.getLogger(__name__)


class AttachmentGenerator:
    """Generate attachments for testing"""
    
    def __init__(self, client: ConfluenceClient, config: Any):
        self.client = client
        self.config = config
        self.temp_dir = Path("/tmp/confluence_attachments")
        self.temp_dir.mkdir(exist_ok=True)
        
    def create_attachment(self, page_id: str, category: str) -> bool:
        """Create and attach a file to a page"""
        file_path = None
        try:
            # Choose attachment type based on category
            attachment_types = {
                'technical': ['pdf', 'png', 'txt'],
                'api_docs': ['pdf', 'json', 'yaml'],
                'knowledge_base': ['pdf', 'docx', 'png'],
                'project_docs': ['xlsx', 'docx', 'pdf'],
                'multilingual': ['pdf', 'docx', 'txt']
            }
            
            file_type = random.choice(
                attachment_types.get(category, ['pdf', 'png', 'txt'])
            )
            
            # Generate file
            file_path = self._generate_file(file_type, category)
            
            if not file_path or not file_path.exists():
                logger.error(f"Failed to generate file of type {file_type}")
                return False
            
            logger.info(f"Generated file: {file_path} (size: {file_path.stat().st_size} bytes)")
            
            # Upload attachment - try both methods
            try:
                # Method 1: attach_file with file path
                result = self.client.atlassian_api.attach_file(
                    filename=str(file_path),
                    page_id=page_id
                )
                logger.info(f"Successfully attached {file_path.name} to page {page_id} using attach_file")
                return True
            except Exception as e:
                logger.warning(f"attach_file failed: {e}, trying attach_content method")
                
                # Method 2: attach_content with file content
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read()
                    
                    # Determine content type
                    content_types = {
                        'pdf': 'application/pdf',
                        'png': 'image/png',
                        'jpg': 'image/jpeg',
                        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                        'txt': 'text/plain',
                        'json': 'application/json',
                        'yaml': 'text/yaml'
                    }
                    file_ext = file_path.suffix[1:].lower()
                    content_type = content_types.get(file_ext, 'application/octet-stream')
                    
                    result = self.client.atlassian_api.attach_content(
                        content=content,
                        name=file_path.name,
                        page_id=page_id,
                        content_type=content_type
                    )
                    logger.info(f"Successfully attached {file_path.name} to page {page_id} using attach_content")
                    return True
                except Exception as e2:
                    logger.error(f"Both attachment methods failed: attach_file: {e}, attach_content: {e2}")
                    raise e2
            
        except Exception as e:
            logger.error(f"Error creating attachment for page {page_id}: {e}", exc_info=True)
            return False
        finally:
            # Clean up
            if file_path and file_path.exists():
                try:
                    file_path.unlink()
                    logger.debug(f"Cleaned up temporary file: {file_path}")
                except Exception as e:
                    logger.warning(f"Failed to clean up file {file_path}: {e}")
    
    def _generate_file(self, file_type: str, category: str) -> Optional[Path]:
        """Generate a file of the specified type"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        generators = {
            'pdf': self._generate_pdf,
            'png': self._generate_image,
            'jpg': self._generate_image,
            'docx': self._generate_docx,
            'xlsx': self._generate_xlsx,
            'txt': self._generate_text,
            'json': self._generate_json,
            'yaml': self._generate_yaml
        }
        
        generator = generators.get(file_type, self._generate_text)
        filename = f"{category}_{timestamp}.{file_type}"
        
        return generator(filename, category)
    
    def _generate_pdf(self, filename: str, category: str) -> Path:
        """Generate a PDF document"""
        file_path = self.temp_dir / filename
        
        c = canvas.Canvas(str(file_path), pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 24)
        c.drawString(100, height - 100, f"{category.title()} Documentation")
        
        # Content
        c.setFont("Helvetica", 12)
        y_position = height - 150
        
        content_lines = [
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"Category: {category}",
            "",
            "This is a sample PDF document generated for testing purposes.",
            "It contains real-looking technical documentation content.",
            "",
            "Key Features:",
            "• Automated testing with real data",
            "• Integration with Confluence API",
            "• Support for multiple content types",
            "",
            "Technical Specifications:",
            "• Format: PDF",
            "• Encoding: UTF-8",
            "• Compatibility: Confluence 7.x+"
        ]
        
        for line in content_lines:
            c.drawString(100, y_position, line)
            y_position -= 20
        
        # Add a simple chart
        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, y_position - 40, "Performance Metrics")
        
        # Draw a simple bar chart
        c.setFillColorRGB(0.2, 0.5, 0.8)
        bars = [
            ("Search", 85),
            ("Index", 92),
            ("Generate", 78),
            ("Export", 88)
        ]
        
        x = 100
        y = y_position - 100
        for label, value in bars:
            c.rect(x, y, 30, value * 2, fill=1)
            c.setFillColorRGB(0, 0, 0)
            c.setFont("Helvetica", 10)
            c.drawString(x, y - 15, label)
            x += 50
        
        c.save()
        return file_path
    
    def _generate_image(self, filename: str, category: str) -> Path:
        """Generate an image file"""
        file_path = self.temp_dir / filename
        
        # Create a figure
        fig, ax = plt.subplots(figsize=(8, 6))
        
        # Generate sample data
        categories = ['Week 1', 'Week 2', 'Week 3', 'Week 4']
        values1 = [random.randint(50, 100) for _ in range(4)]
        values2 = [random.randint(30, 80) for _ in range(4)]
        
        x = range(len(categories))
        width = 0.35
        
        ax.bar([i - width/2 for i in x], values1, width, label='Metric A')
        ax.bar([i + width/2 for i in x], values2, width, label='Metric B')
        
        ax.set_xlabel('Time Period')
        ax.set_ylabel('Performance Score')
        ax.set_title(f'{category.title()} Performance Metrics')
        ax.set_xticks(x)
        ax.set_xticklabels(categories)
        ax.legend()
        ax.grid(True, alpha=0.3)
        
        plt.tight_layout()
        plt.savefig(file_path, dpi=150)
        plt.close()
        
        return file_path
    
    def _generate_docx(self, filename: str, category: str) -> Path:
        """Generate a Word document"""
        file_path = self.temp_dir / filename
        
        doc = Document()
        
        # Title
        doc.add_heading(f'{category.title()} Technical Report', 0)
        
        # Metadata
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'Category: {category}')
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(
            'This document provides a comprehensive overview of the technical '
            'implementation and testing procedures for the Confluence Gateway system. '
            'The content has been generated using real data collection methods.'
        )
        
        # Technical Details
        doc.add_heading('Technical Details', level=1)
        doc.add_paragraph('The system architecture consists of:')
        
        # Add a list
        doc.add_paragraph('• Service Layer Architecture', style='List Bullet')
        doc.add_paragraph('• Adapter Pattern Implementation', style='List Bullet')
        doc.add_paragraph('• Real Data Collection Pipeline', style='List Bullet')
        doc.add_paragraph('• Quality Assurance Metrics', style='List Bullet')
        
        # Code Example
        doc.add_heading('Code Example', level=2)
        doc.add_paragraph('Sample configuration:', style='Intense Quote')
        
        code = doc.add_paragraph()
        code.add_run('''
{
    "collection_settings": {
        "max_content_size_mb": 10,
        "quality_threshold": 0.7,
        "rate_limit_seconds": 2
    }
}
        ''').font.name = 'Courier New'
        
        # Conclusion
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph(
            'The implementation successfully demonstrates the capability to collect '
            'and process real documentation for testing purposes.'
        )
        
        doc.save(file_path)
        return file_path
    
    def _generate_xlsx(self, filename: str, category: str) -> Path:
        """Generate an Excel spreadsheet"""
        file_path = self.temp_dir / filename
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Metrics"
        
        # Headers
        headers = ['Date', 'Category', 'Pages Created', 'Quality Score', 'Processing Time (s)']
        ws.append(headers)
        
        # Sample data
        for i in range(20):
            date = datetime.now().strftime('%Y-%m-%d')
            row = [
                date,
                random.choice(['technical', 'api_docs', 'knowledge_base']),
                random.randint(10, 100),
                round(random.uniform(0.7, 1.0), 2),
                round(random.uniform(0.5, 5.0), 2)
            ]
            ws.append(row)
        
        # Add summary sheet
        summary = wb.create_sheet("Summary")
        summary.append(['Metric', 'Value'])
        summary.append(['Total Pages', '=SUM(Metrics!C2:C21)'])
        summary.append(['Average Quality', '=AVERAGE(Metrics!D2:D21)'])
        summary.append(['Average Time', '=AVERAGE(Metrics!E2:E21)'])
        
        # Style the headers
        for cell in ws[1]:
            cell.font = cell.font.copy(bold=True)
        for cell in summary[1]:
            cell.font = cell.font.copy(bold=True)
        
        wb.save(file_path)
        return file_path
    
    def _generate_text(self, filename: str, category: str) -> Path:
        """Generate a text file"""
        file_path = self.temp_dir / filename
        
        content = f"""
{category.upper()} DOCUMENTATION
{'=' * 50}

Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Category: {category}
Format: Plain Text

OVERVIEW
--------
This file contains sample documentation content for testing the Confluence
Gateway system with real data. The content is structured to simulate actual
technical documentation.

KEY CONCEPTS
------------
1. Real Data Collection
   - Automated scraping from public sources
   - GitHub repository documentation
   - Web-based technical documentation

2. Content Processing
   - Format conversion (Markdown, HTML, RST)
   - Quality scoring and validation
   - Metadata extraction

3. Integration Testing
   - Confluence API compatibility
   - Search functionality verification
   - Performance benchmarking

IMPLEMENTATION NOTES
-------------------
The system uses a modular architecture with the following components:

- ContentCollector: Fetches content from various sources
- ContentProcessor: Cleans and formats content
- QualityAnalyzer: Scores content quality
- ConfluenceUploader: Manages page creation and updates

CONFIGURATION
-------------
Sample configuration parameters:

max_content_size_mb: 10
min_content_length: 500
quality_threshold: 0.7
rate_limit_seconds: 2

TESTING PROCEDURES
-----------------
1. Run real_data_collector.py to gather content
2. Execute generate_dummy_data.py with --real-data flag
3. Verify content in Confluence
4. Run search tests to validate indexing

CONCLUSION
----------
This documentation demonstrates the capability to work with real content
for comprehensive testing of the Confluence Gateway system.
"""
        
        with open(file_path, 'w') as f:
            f.write(content)
        
        return file_path
    
    def _generate_json(self, filename: str, category: str) -> Path:
        """Generate a JSON file"""
        file_path = self.temp_dir / filename
        
        data = {
            "metadata": {
                "generated": datetime.now().isoformat(),
                "category": category,
                "version": "1.0.0"
            },
            "configuration": {
                "api_endpoints": [
                    "/api/v1/search",
                    "/api/v1/index",
                    "/api/v1/generate"
                ],
                "supported_formats": ["markdown", "html", "rst", "plain_text"],
                "quality_thresholds": {
                    "minimum": 0.5,
                    "recommended": 0.7,
                    "optimal": 0.9
                }
            },
            "sample_data": {
                "pages_processed": random.randint(100, 1000),
                "average_quality": round(random.uniform(0.7, 0.95), 3),
                "categories": {
                    "technical": random.randint(20, 100),
                    "api_docs": random.randint(15, 80),
                    "knowledge_base": random.randint(25, 90)
                }
            }
        }
        
        with open(file_path, 'w') as f:
            json.dump(data, f, indent=2)
        
        return file_path
    
    def _generate_yaml(self, filename: str, category: str) -> Path:
        """Generate a YAML file"""
        file_path = self.temp_dir / filename
        
        content = f"""# {category.title()} Configuration
# Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

metadata:
  version: 1.0.0
  category: {category}
  environment: testing

api_configuration:
  base_url: https://confluence.example.com
  endpoints:
    - /rest/api/content
    - /rest/api/search
    - /rest/api/space
  
search_settings:
  algorithms:
    - keyword
    - semantic
    - hybrid
  weights:
    keyword: 0.4
    semantic: 0.6
  
quality_metrics:
  thresholds:
    low: 0.5
    medium: 0.7
    high: 0.9
  factors:
    - readability
    - technical_depth
    - completeness
    - accuracy

performance_targets:
  search_latency_ms: 100
  index_throughput_docs_per_sec: 50
  generation_time_sec: 2.0
"""
        
        with open(file_path, 'w') as f:
            f.write(content)
        
        return file_path